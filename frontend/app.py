import streamlit as st
import pandas as pd
from docx import Document
import os
from dotenv import load_dotenv
import google.generativeai as genai

# Configurações
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=GOOGLE_API_KEY)

# Modelo atualizado (usando o mais recente disponível)
MODEL_NAME = "gemini-1.5-flash"  # Modelo rápido e gratuito (ou "gemini-1.5-pro" para melhor qualidade)

# Título da aplicação
st.title("📝 Gerador de Relatórios Ambientais")
st.markdown("Transforme planilhas em relatórios técnicos usando IA do Google!")

# Upload da planilha
uploaded_file = st.file_uploader("**Envie sua planilha** (Excel ou CSV)", type=["xlsx", "csv"])

# Inputs do usuário
report_type = st.selectbox(
    "**Tipo de relatório**",
    ("Impactos Hidrológicos", "Qualidade do Ar", "Gestão de Resíduos")
)

# Processamento
if uploaded_file and st.button("Gerar Relatório"):
    with st.spinner("Processando"):
        try:
            # 1. Lê os dados (limita a 100 linhas para evitar erros)
            df = pd.read_excel(uploaded_file) if uploaded_file.name.endswith('.xlsx') else pd.read_csv(uploaded_file)
            sample_data = df.head(100).to_string()  # Amostra para o prompt
            
            # 2. Configurações de geração
            generation_config = {
                "temperature": 0.3,  # Mais determinístico (ideal para relatórios técnicos)
                "max_output_tokens": 2000  # Limite de tokens
            }
            
            safety_settings = {
                "HARM_CATEGORY_HARASSMENT": "BLOCK_NONE",
                "HARM_CATEGORY_HATE_SPEECH": "BLOCK_NONE",
                "HARM_CATEGORY_SEXUALLY_EXPLICIT": "BLOCK_NONE",
                "HARM_CATEGORY_DANGEROUS_CONTENT": "BLOCK_NONE"
            }
            
            # 3. Gera o relatório
            model = genai.GenerativeModel(MODEL_NAME)
            prompt = f"""
            Gere um relatório técnico de {report_type} para engenheiros ambientais com base nos dados abaixo.
            
            Requisitos:
            - Formato profissional com introdução, análise e conclusões
            - Cite normas brasileiras (ex: CONAMA 357/05, NBR 10004) quando relevante
            - Seja conciso e técnico
            - Dados fornecidos (amostra):
            {sample_data}
            """
            
            response = model.generate_content(
                contents=[prompt],
                generation_config=generation_config,
                safety_settings=safety_settings
            )
            
            if not response.text:
                raise Exception("Resposta vazia - verifique seu prompt ou dados")
                
            report_text = response.text

            # 4. Cria documento Word
            doc = Document()
            doc.add_heading(f"Relatório de {report_type} (Gerado por Gemini 1.5)", level=1)
            for paragraph in report_text.split('\n'):
                doc.add_paragraph(paragraph)
            doc.save("relatorio_gemini.docx")
            
            # 5. Exibe resultados
            st.success("✅ Relatório gerado com sucesso!")
            st.download_button(
                label="⬇️ Baixar DOCX",
                data=open("relatorio_gemini.docx", "rb").read(),
                file_name=f"relatorio_{report_type.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            with st.expander("📄 Visualizar texto completo"):
                st.write(report_text)
                
        except Exception as e:
            st.error(f"❌ Erro: {str(e)}")
            st.info("Dicas:\n1. Verifique sua API Key\n2. Reduza o tamanho dos dados de entrada\n3. Tente novamente mais tarde")